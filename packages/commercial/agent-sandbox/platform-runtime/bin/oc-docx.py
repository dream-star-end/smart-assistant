#!/usr/bin/env python3
"""OpenClaude Word/DOCX generation, rendering, inspection, and privacy cleanup.

The structured builder and QA workflow are adapted from the user-supplied
Word DOCX Skill Kit v1.0.0. The original kit is distributed under the MIT
License:

Copyright (c) 2026

Permission is hereby granted, free of charge, to any person obtaining a copy
of this software and associated documentation files (the "Software"), to deal
in the Software without restriction, including without limitation the rights
to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
copies of the Software, and to permit persons to whom the Software is
furnished to do so, subject to the following conditions:

The above copyright notice and this permission notice shall be included in all
copies or substantial portions of the Software.

THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
SOFTWARE.
"""

from __future__ import annotations

import argparse
import io
import json
import os
import shutil
import subprocess
import sys
import tempfile
import zipfile
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Any, Iterable


DEFAULT_REFERENCE_DOC = Path("/usr/local/share/openclaude-docgen/reference.docx")
VISION_COPY_MAX_BYTES = 4_500_000
VISION_COPY_MAX_EDGE = 1600
# 本 CLI 无 sibling 引用，SELF_ROOT 仅立“工具单文件独立、禁相对 sibling 裸调用”不变量（测试固化）。
SELF_ROOT = Path(__file__).resolve().parent.parent


class ConfigError(ValueError):
    """Raised when a structured document configuration is invalid."""


class RenderError(RuntimeError):
    """Raised when DOCX rendering fails."""


# Structured authoring dependencies are intentionally lazy. `convert` and all
# `--help` surfaces keep working in source-only CI probes; the runtime image
# installs the full dependency set and its image smoke exercises every command.
Document = None
WD_STYLE_TYPE = WD_CELL_VERTICAL_ALIGNMENT = None
WD_TABLE_ALIGNMENT = WD_ALIGN_PARAGRAPH = WD_LINE_SPACING = None
OxmlElement = qn = Cm = Pt = RGBColor = None
Image = None
etree = None
yaml = None


def _load_docx_dependencies() -> None:
    global Document, WD_STYLE_TYPE, WD_CELL_VERTICAL_ALIGNMENT
    global WD_TABLE_ALIGNMENT, WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    global OxmlElement, qn, Cm, Pt, RGBColor
    if Document is not None:
        return
    try:
        from docx import Document as _Document
        from docx.enum.style import WD_STYLE_TYPE as _WD_STYLE_TYPE
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT as _WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.table import WD_TABLE_ALIGNMENT as _WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD_ALIGN_PARAGRAPH
        from docx.enum.text import WD_LINE_SPACING as _WD_LINE_SPACING
        from docx.oxml import OxmlElement as _OxmlElement
        from docx.oxml.ns import qn as _qn
        from docx.shared import Cm as _Cm
        from docx.shared import Pt as _Pt
        from docx.shared import RGBColor as _RGBColor
    except ImportError as exc:
        raise RuntimeError("缺少 python-docx；请使用 OpenClaude runtime image") from exc
    Document = _Document
    WD_STYLE_TYPE = _WD_STYLE_TYPE
    WD_CELL_VERTICAL_ALIGNMENT = _WD_CELL_VERTICAL_ALIGNMENT
    WD_TABLE_ALIGNMENT = _WD_TABLE_ALIGNMENT
    WD_ALIGN_PARAGRAPH = _WD_ALIGN_PARAGRAPH
    WD_LINE_SPACING = _WD_LINE_SPACING
    OxmlElement = _OxmlElement
    qn = _qn
    Cm = _Cm
    Pt = _Pt
    RGBColor = _RGBColor


def _load_yaml_dependency() -> None:
    global yaml
    if yaml is not None:
        return
    try:
        import yaml as _yaml
    except ImportError as exc:
        raise RuntimeError("缺少 PyYAML；请使用 OpenClaude runtime image") from exc
    yaml = _yaml


def _load_pillow_dependency() -> None:
    global Image
    if Image is not None:
        return
    try:
        from PIL import Image as _Image
    except ImportError as exc:
        raise RuntimeError("缺少 Pillow；请使用 OpenClaude runtime image") from exc
    Image = _Image


def _load_lxml_dependency() -> None:
    global etree
    if etree is not None:
        return
    try:
        from lxml import etree as _etree
    except ImportError as exc:
        raise RuntimeError("缺少 lxml；请使用 OpenClaude runtime image") from exc
    etree = _etree


def _natural_page_number(path: Path) -> int:
    try:
        return int(path.stem.rsplit("-", 1)[-1])
    except ValueError as exc:
        raise ValueError(f"无法解析页面编号：{path.name}") from exc


def _which(*names: str) -> str:
    for name in names:
        path = shutil.which(name)
        if path:
            return path
    raise RenderError(f"找不到可执行程序：{', '.join(names)}")


def _resolve_existing_file(value: str | Path, label: str) -> Path:
    path = Path(value).expanduser().resolve()
    if not path.is_file():
        raise FileNotFoundError(f"{label}不存在：{path}")
    return path


def _resolve_output(value: str | Path) -> Path:
    output = Path(value).expanduser().resolve()
    output.parent.mkdir(parents=True, exist_ok=True)
    return output


def convert_docx(
    input_path: str | Path,
    output_path: str | Path | None,
    *,
    engine: str | None,
    reference_doc: str | Path | None,
    use_reference: bool,
) -> Path:
    source = _resolve_existing_file(input_path, "输入文件")
    extension = source.suffix.lower()
    selected_engine = engine or ("quarto" if extension == ".qmd" else "pandoc")
    output = _resolve_output(output_path or source.with_suffix(".docx"))

    reference: Path | None = None
    if use_reference:
        reference = _resolve_existing_file(reference_doc or DEFAULT_REFERENCE_DOC, "reference doc")

    if selected_engine == "pandoc":
        command = [
            _which("pandoc"),
            str(source),
            "--standalone",
            "--from=markdown+tex_math_dollars+pipe_tables+raw_html+fenced_divs+footnotes+definition_lists",
            "--to=docx",
            "--output",
            str(output),
            "--toc",
            "--number-sections",
        ]
        if reference:
            command.extend(["--reference-doc", str(reference)])
        subprocess.run(command, check=True)
    elif selected_engine == "quarto":
        render_name = output.name
        rendered = source.parent / render_name
        if rendered != output:
            render_name = f".oc-docx-{source.stem}-{os.getpid()}.docx"
            rendered = source.parent / render_name
        rendered.unlink(missing_ok=True)
        command = [_which("quarto"), "render", source.name, "--to", "docx", "--output", render_name]
        if reference:
            command.extend(["-M", f"reference-doc={reference}"])
        try:
            subprocess.run(command, cwd=source.parent, check=True)
            if rendered != output:
                os.replace(rendered, output)
        finally:
            if rendered != output:
                rendered.unlink(missing_ok=True)
    else:
        raise ValueError(f"无效引擎：{selected_engine}")

    if not output.is_file() or output.stat().st_size == 0:
        raise RuntimeError(f"DOCX 未生成：{output}")
    return output


def load_config(path: str | Path) -> dict[str, Any]:
    source = _resolve_existing_file(path, "配置文件")
    suffix = source.suffix.lower()
    with source.open("r", encoding="utf-8") as handle:
        if suffix in {".yaml", ".yml"}:
            _load_yaml_dependency()
            data = yaml.safe_load(handle)
        elif suffix == ".json":
            data = json.load(handle)
        else:
            raise ConfigError("build 仅支持 .yaml、.yml 或 .json 配置")
    if not isinstance(data, dict):
        raise ConfigError("配置文件顶层必须是对象")
    if not isinstance(data.get("document"), dict):
        raise ConfigError("缺少 document 对象")
    if not isinstance(data.get("sections"), list):
        raise ConfigError("缺少 sections 数组")
    if not str(data["document"].get("title", "")).strip():
        raise ConfigError("document.title 不能为空")
    return data


def apply_overrides(config: dict[str, Any], overrides: list[str] | None) -> dict[str, Any]:
    for item in overrides or []:
        if "=" not in item:
            raise ConfigError(f"覆盖参数必须为 key=value：{item}")
        dotted_key, raw_value = item.split("=", 1)
        keys = [part for part in dotted_key.split(".") if part]
        if not keys:
            raise ConfigError(f"无效覆盖键：{item}")
        lowered = raw_value.lower()
        value: Any
        if lowered in {"true", "false"}:
            value = lowered == "true"
        elif lowered in {"null", "none"}:
            value = None
        else:
            try:
                value = json.loads(raw_value)
            except json.JSONDecodeError:
                value = raw_value
        target: dict[str, Any] = config
        for key in keys[:-1]:
            existing = target.get(key)
            if not isinstance(existing, dict):
                existing = {}
                target[key] = existing
            target = existing
        target[keys[-1]] = value
    return config


@dataclass(frozen=True)
class FontSettings:
    east_asia: str = "Noto Sans CJK SC"
    latin: str = "DejaVu Sans"
    mono: str = "DejaVu Sans Mono"
    body_size: float = 10.5


@dataclass(frozen=True)
class Theme:
    accent: str = "1F4E78"
    accent_light: str = "D9EAF7"
    heading: str = "17365D"
    border: str = "B4C6E7"
    code_fill: str = "F3F5F7"
    muted: str = "666666"


def _hex_color(value: Any, default: str) -> str:
    candidate = str(value or "").replace("#", "").strip().upper()
    if len(candidate) == 6 and all(character in "0123456789ABCDEF" for character in candidate):
        return candidate
    return default


def _set_run_font(run: Any, east_asia: str, latin: str, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = latin
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def _set_style_font(style: Any, east_asia: str, latin: str, size: float | None = None, bold: bool | None = None) -> None:
    style.font.name = latin
    fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    if size is not None:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill.replace("#", "").upper())


def set_cell_margins(cell: Any, top: int = 90, start: int = 110, bottom: int = 90, end: int = 110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def _toggle_paragraph_property(paragraph: Any, tag: str, enabled: bool = True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    element = p_pr.find(qn(f"w:{tag}"))
    if enabled and element is None:
        p_pr.append(OxmlElement(f"w:{tag}"))
    elif not enabled and element is not None:
        p_pr.remove(element)


def set_keep_with_next(paragraph: Any, enabled: bool = True) -> None:
    _toggle_paragraph_property(paragraph, "keepNext", enabled)


def set_keep_lines(paragraph: Any, enabled: bool = True) -> None:
    _toggle_paragraph_property(paragraph, "keepLines", enabled)


def add_field(paragraph: Any, field_code: str, placeholder: str = "") -> None:
    begin_run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run.append(begin)
    instruction_run = OxmlElement("w:r")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = field_code
    instruction_run.append(instruction)
    separate_run = OxmlElement("w:r")
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run.append(separate)
    display_run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = placeholder
    display_run.append(text)
    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)
    paragraph._p.extend([begin_run, instruction_run, separate_run, display_run, end_run])


def set_paragraph_borders(paragraph: Any, color: str = "D9E2F3", size: int = 8, space: int = 3) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for side in ("top", "left", "bottom", "right"):
        edge = borders.find(qn(f"w:{side}"))
        if edge is None:
            edge = OxmlElement(f"w:{side}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:space"), str(space))
        edge.set(qn("w:color"), color.replace("#", "").upper())


def request_field_update(document: Any) -> None:
    settings = document.settings._element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def _set_cell_text(cell: Any, text: Any, font: FontSettings, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("" if text is None else str(text))
    _set_run_font(run, font.east_asia, font.latin, font.body_size, bold)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def _add_hyperlink(paragraph: Any, text: str, url: str, font: FontSettings) -> None:
    relation_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    new_run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:eastAsia"), font.east_asia)
    fonts.set(qn("w:ascii"), font.latin)
    fonts.set(qn("w:hAnsi"), font.latin)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend([fonts, color, underline])
    new_run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


class DocumentBuilder:
    def __init__(self, config: dict[str, Any]):
        _load_docx_dependencies()
        self.config = config
        document_config = config.get("document", {})
        font_config = document_config.get("fonts", {}) or {}
        theme_config = document_config.get("theme", {}) or {}
        self.font = FontSettings(
            east_asia=str(font_config.get("east_asia") or "Noto Sans CJK SC"),
            latin=str(font_config.get("latin") or "DejaVu Sans"),
            mono=str(font_config.get("mono") or "DejaVu Sans Mono"),
            body_size=float(font_config.get("body_size") or 10.5),
        )
        self.theme = Theme(
            accent=_hex_color(theme_config.get("accent"), "1F4E78"),
            accent_light=_hex_color(theme_config.get("accent_light"), "D9EAF7"),
            heading=_hex_color(theme_config.get("heading"), "17365D"),
            border=_hex_color(theme_config.get("border"), "B4C6E7"),
            code_fill=_hex_color(theme_config.get("code_fill"), "F3F5F7"),
            muted=_hex_color(theme_config.get("muted"), "666666"),
        )
        self.document = Document()
        self._configure_page()
        self._configure_styles()
        self._configure_header_footer()

    def _configure_page(self) -> None:
        document_config = self.config.get("document", {})
        page_config = document_config.get("page", {}) or {}
        section = self.document.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(float(page_config.get("top_cm", 2.2)))
        section.bottom_margin = Cm(float(page_config.get("bottom_cm", 2.0)))
        section.left_margin = Cm(float(page_config.get("left_cm", 2.4)))
        section.right_margin = Cm(float(page_config.get("right_cm", 2.2)))
        section.header_distance = Cm(float(page_config.get("header_cm", 1.0)))
        section.footer_distance = Cm(float(page_config.get("footer_cm", 1.0)))

    def _configure_styles(self) -> None:
        styles = self.document.styles
        normal = styles["Normal"]
        _set_style_font(normal, self.font.east_asia, self.font.latin, self.font.body_size)
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        normal.paragraph_format.line_spacing = 1.35
        normal.paragraph_format.space_after = Pt(6)
        title = styles["Title"]
        _set_style_font(title, self.font.east_asia, self.font.latin, 26, True)
        title.font.color.rgb = RGBColor.from_string(self.theme.heading)
        title.paragraph_format.space_after = Pt(12)
        subtitle = styles["Subtitle"]
        _set_style_font(subtitle, self.font.east_asia, self.font.latin, 13)
        subtitle.font.color.rgb = RGBColor.from_string(self.theme.muted)
        for level, size in {1: 18, 2: 14, 3: 12}.items():
            style = styles[f"Heading {level}"]
            _set_style_font(style, self.font.east_asia, self.font.latin, size, True)
            style.font.color.rgb = RGBColor.from_string(self.theme.heading)
            style.paragraph_format.space_before = Pt(12 if level == 1 else 9)
            style.paragraph_format.space_after = Pt(5)
            style.paragraph_format.keep_with_next = True
            style.paragraph_format.keep_together = True
        code = styles["Code Block"] if "Code Block" in styles else styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        _set_style_font(code, self.font.east_asia, self.font.mono, 9.5)
        code.paragraph_format.left_indent = Cm(0.4)
        code.paragraph_format.right_indent = Cm(0.2)
        code.paragraph_format.space_before = Pt(4)
        code.paragraph_format.space_after = Pt(7)
        code.paragraph_format.line_spacing = 1.05
        quote = styles["Quote"] if "Quote" in styles else styles.add_style("Quote", WD_STYLE_TYPE.PARAGRAPH)
        _set_style_font(quote, self.font.east_asia, self.font.latin, 10.5)
        quote.font.italic = True
        quote.font.color.rgb = RGBColor.from_string(self.theme.muted)
        quote.paragraph_format.left_indent = Cm(0.7)
        quote.paragraph_format.right_indent = Cm(0.4)
        quote.paragraph_format.space_before = Pt(4)
        quote.paragraph_format.space_after = Pt(7)
        for style_name in ("List Bullet", "List Number"):
            if style_name in styles:
                _set_style_font(styles[style_name], self.font.east_asia, self.font.latin, self.font.body_size)

    def _configure_header_footer(self) -> None:
        document_config = self.config.get("document", {})
        title = str(document_config.get("short_title") or document_config.get("title") or "")
        header_text = str(document_config.get("header") or title)
        footer_text = str(document_config.get("footer") or "")
        for section in self.document.sections:
            if bool(document_config.get("cover", True)):
                section.different_first_page_header_footer = True
                section.first_page_header.paragraphs[0].text = ""
                section.first_page_footer.paragraphs[0].text = ""
            header = section.header.paragraphs[0]
            header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = header.add_run(header_text)
            _set_run_font(run, self.font.east_asia, self.font.latin, 8.5)
            run.font.color.rgb = RGBColor.from_string(self.theme.muted)
            footer = section.footer.paragraphs[0]
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if footer_text:
                run = footer.add_run(f"{footer_text}  ·  ")
                _set_run_font(run, self.font.east_asia, self.font.latin, 8.5)
                run.font.color.rgb = RGBColor.from_string(self.theme.muted)
            add_field(footer, " PAGE ", "1")

    def _add_cover(self) -> None:
        config = self.config.get("document", {})
        spacer = self.document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(55)
        title = self.document.add_paragraph(style="Title")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_font(title.add_run(str(config.get("title", ""))), self.font.east_asia, self.font.latin, 26, True)
        subtitle_text = str(config.get("subtitle", ""))
        if subtitle_text:
            subtitle = self.document.add_paragraph(style="Subtitle")
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_run_font(subtitle.add_run(subtitle_text), self.font.east_asia, self.font.latin, 13)
        line = self.document.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line.paragraph_format.space_before = Pt(14)
        run = line.add_run("━" * 18)
        _set_run_font(run, self.font.east_asia, self.font.latin, 10)
        run.font.color.rgb = RGBColor.from_string(self.theme.accent)
        metadata = [
            str(config.get("author") or ""),
            f"版本：{config['version']}" if config.get("version") else "",
            str(config.get("date") or date.today().isoformat()),
        ]
        for value in [item for item in metadata if item]:
            paragraph = self.document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(3)
            run = paragraph.add_run(value)
            _set_run_font(run, self.font.east_asia, self.font.latin, 10.5)
            run.font.color.rgb = RGBColor.from_string(self.theme.muted)
        self.document.add_page_break()

    def _add_toc(self) -> None:
        config = self.config.get("document", {})
        mode = str(config.get("table_of_contents_mode") or "static").lower()
        if mode not in {"static", "field", "both"}:
            raise ValueError("document.table_of_contents_mode 仅支持 static、field 或 both")
        heading = self.document.add_heading("目录", level=1)
        set_keep_with_next(heading)
        if mode in {"field", "both"}:
            add_field(self.document.add_paragraph(), ' TOC \\o "1-3" \\h \\z \\u ', "目录将在 Word 中更新")
        if mode in {"static", "both"}:
            for section in self.config.get("sections", []):
                if not isinstance(section, dict):
                    continue
                text = str(section.get("heading") or "").strip()
                if not text:
                    continue
                level = min(max(int(section.get("level", 1)), 1), 3)
                paragraph = self.document.add_paragraph()
                paragraph.paragraph_format.left_indent = Cm(0.55 * (level - 1))
                paragraph.paragraph_format.space_after = Pt(5)
                run = paragraph.add_run(text)
                _set_run_font(run, self.font.east_asia, self.font.latin, 11 if level == 1 else 10.5, level == 1)
                if level == 1:
                    run.font.color.rgb = RGBColor.from_string(self.theme.heading)
        notes = {
            "field": "提示：首次在 Microsoft Word 中打开时，可按 Ctrl+A、F9 更新目录和页码。",
            "both": "提示：上方包含自动目录字段和静态导航；在 Microsoft Word 中按 Ctrl+A、F9 可更新自动目录。",
            "static": "提示：当前为跨渲染器稳定显示的静态内容导航，不包含自动页码。",
        }
        note = self.document.add_paragraph(notes[mode])
        note.style = self.document.styles["Quote"]
        self.document.add_page_break()

    def _add_heading(self, text: str, level: int) -> None:
        level = min(max(int(level), 1), 3)
        paragraph = self.document.add_heading(text, level=level)
        set_keep_with_next(paragraph)
        set_keep_lines(paragraph)
        if level == 1:
            set_paragraph_borders(paragraph, color=self.theme.border, size=6, space=2)

    def _add_paragraph(self, block: dict[str, Any]) -> None:
        style = str(block.get("style") or "Normal")
        paragraph = self.document.add_paragraph(style=style if style in self.document.styles else "Normal")
        paragraph.alignment = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }.get(str(block.get("align", "justify")).lower(), WD_ALIGN_PARAGRAPH.JUSTIFY)
        paragraph.paragraph_format.first_line_indent = (
            Cm(float(block.get("first_line_cm", 0.74))) if block.get("first_line", True) else None
        )
        run = paragraph.add_run(str(block.get("text", "")))
        _set_run_font(run, self.font.east_asia, self.font.latin, float(block.get("font_size", self.font.body_size)))
        run.bold = bool(block.get("bold"))
        run.italic = bool(block.get("italic"))
        if block.get("keep_together"):
            set_keep_lines(paragraph)

    def _add_list(self, block: dict[str, Any], numbered: bool) -> None:
        items = block.get("items", [])
        if not isinstance(items, list):
            raise ValueError("list.items 必须是数组")
        for item in items:
            paragraph = self.document.add_paragraph(style="List Number" if numbered else "List Bullet")
            paragraph.paragraph_format.space_after = Pt(2)
            _set_run_font(paragraph.add_run(str(item)), self.font.east_asia, self.font.latin, self.font.body_size)

    def _add_code(self, block: dict[str, Any]) -> None:
        language = str(block.get("language") or "").strip()
        if language:
            label = self.document.add_paragraph()
            label.paragraph_format.space_after = Pt(2)
            run = label.add_run(language.upper())
            _set_run_font(run, self.font.east_asia, self.font.latin, 8.5, True)
            run.font.color.rgb = RGBColor.from_string(self.theme.muted)
            set_keep_with_next(label)
        paragraph = self.document.add_paragraph(style="Code Block")
        run = paragraph.add_run(str(block.get("text", "")).rstrip("\n"))
        _set_run_font(run, self.font.east_asia, self.font.mono, float(block.get("font_size", 9.5)))
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), self.theme.code_fill)
        paragraph._p.get_or_add_pPr().append(shading)
        set_paragraph_borders(paragraph, color=self.theme.border, size=4, space=3)
        set_keep_lines(paragraph)

    def _add_quote(self, block: dict[str, Any]) -> None:
        paragraph = self.document.add_paragraph(style="Quote")
        _set_run_font(paragraph.add_run(str(block.get("text", ""))), self.font.east_asia, self.font.latin, self.font.body_size)
        set_paragraph_borders(paragraph, color=self.theme.accent, size=10, space=4)

    def _add_table(self, block: dict[str, Any]) -> None:
        headers = block.get("headers", [])
        rows = block.get("rows", [])
        if not isinstance(headers, list) or not headers:
            raise ValueError("table.headers 必须是非空数组")
        if not isinstance(rows, list):
            raise ValueError("table.rows 必须是数组")
        caption = str(block.get("caption") or "").strip()
        if caption:
            paragraph = self.document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(caption)
            _set_run_font(run, self.font.east_asia, self.font.latin, 9.5, True)
            run.font.color.rgb = RGBColor.from_string(self.theme.muted)
            set_keep_with_next(paragraph)
        table = self.document.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        table.autofit = True
        set_repeat_table_header(table.rows[0])
        for index, header in enumerate(headers):
            _set_cell_text(table.rows[0].cells[index], header, self.font, bold=True, color="FFFFFF")
            set_cell_shading(table.rows[0].cells[index], self.theme.accent)
        for row_index, row_data in enumerate(rows):
            values = row_data if isinstance(row_data, list) else [row_data]
            row = table.add_row()
            for index in range(len(headers)):
                _set_cell_text(row.cells[index], values[index] if index < len(values) else "", self.font)
                if row_index % 2 == 1:
                    set_cell_shading(row.cells[index], "F8FAFC")
        self.document.add_paragraph().paragraph_format.space_after = Pt(2)

    def _add_callout(self, block: dict[str, Any]) -> None:
        kind = str(block.get("kind") or "info").lower()
        fill_map = {"info": "EAF3F8", "warning": "FFF4CE", "success": "E2F0D9", "danger": "FCE4D6"}
        title_map = {"info": "说明", "warning": "注意", "success": "结论", "danger": "风险"}
        table = self.document.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        set_cell_shading(cell, fill_map.get(kind, fill_map["info"]))
        set_cell_margins(cell, top=130, start=170, bottom=130, end=170)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(2)
        title = str(block.get("title") or title_map.get(kind, "说明"))
        _set_run_font(paragraph.add_run(f"{title}："), self.font.east_asia, self.font.latin, self.font.body_size, True)
        _set_run_font(paragraph.add_run(str(block.get("text", ""))), self.font.east_asia, self.font.latin, self.font.body_size)
        self.document.add_paragraph().paragraph_format.space_after = Pt(1)

    def _add_link(self, block: dict[str, Any]) -> None:
        paragraph = self.document.add_paragraph()
        prefix = str(block.get("prefix") or "")
        if prefix:
            _set_run_font(paragraph.add_run(prefix), self.font.east_asia, self.font.latin, self.font.body_size)
        url = str(block.get("url") or "")
        _add_hyperlink(paragraph, str(block.get("label") or url or "链接"), url, self.font)

    def _add_blocks(self, blocks: Iterable[dict[str, Any]]) -> None:
        for block in blocks:
            if not isinstance(block, dict):
                raise ValueError("section.blocks 的每一项必须是对象")
            kind = str(block.get("type") or "paragraph").lower()
            handlers = {
                "paragraph": self._add_paragraph,
                "code": self._add_code,
                "quote": self._add_quote,
                "table": self._add_table,
                "callout": self._add_callout,
                "link": self._add_link,
            }
            if kind == "bullets":
                self._add_list(block, numbered=False)
            elif kind == "numbered":
                self._add_list(block, numbered=True)
            elif kind == "page_break":
                self.document.add_page_break()
            elif kind in handlers:
                handlers[kind](block)
            else:
                raise ValueError(f"未知 block.type：{kind}")

    def build(self, output_path: str | Path, *, cover: bool | None = None, toc: bool | None = None) -> Path:
        config = self.config.get("document", {})
        use_cover = bool(config.get("cover", True)) if cover is None else cover
        use_toc = bool(config.get("table_of_contents", False)) if toc is None else toc
        if use_cover:
            self._add_cover()
        if use_toc:
            self._add_toc()
        for section in self.config.get("sections", []):
            if not isinstance(section, dict):
                raise ValueError("sections 的每一项必须是对象")
            heading = str(section.get("heading") or "").strip()
            if heading:
                self._add_heading(heading, int(section.get("level", 1)))
            blocks = section.get("blocks", [])
            if not isinstance(blocks, list):
                raise ValueError("section.blocks 必须是数组")
            self._add_blocks(blocks)
        core = self.document.core_properties
        core.title = str(config.get("title") or "")
        core.subject = str(config.get("subject") or "")
        core.author = str(config.get("author") or "")
        core.comments = ""
        core.keywords = str(config.get("keywords") or "")
        request_field_update(self.document)
        output = _resolve_output(output_path)
        self.document.save(output)
        return output


def _make_vision_copy(source: Path, destination: Path) -> dict[str, Any]:
    _load_pillow_dependency()
    with Image.open(source) as opened:
        original = opened.convert("RGB")
        max_edge = VISION_COPY_MAX_EDGE
        quality = 85
        while True:
            candidate = original.copy()
            candidate.thumbnail((max_edge, max_edge), Image.Resampling.LANCZOS)
            buffer = io.BytesIO()
            candidate.save(buffer, format="JPEG", quality=quality, optimize=True, progressive=True)
            payload = buffer.getvalue()
            if len(payload) < VISION_COPY_MAX_BYTES:
                destination.write_bytes(payload)
                return {
                    "file": destination.name,
                    "width": candidate.width,
                    "height": candidate.height,
                    "size_bytes": len(payload),
                }
            if quality > 55:
                quality -= 10
            elif max_edge > 320:
                max_edge = max(320, int(max_edge * 0.8))
                quality = 85
            else:
                raise RenderError(f"无法把视觉质检副本压缩到 4.5MB 以下：{source.name}")


def render_docx(
    input_path: str | Path,
    output_dir: str | Path,
    *,
    emit_pdf: bool = False,
    dpi: int = 160,
    verbose: bool = False,
) -> dict[str, Any]:
    source = _resolve_existing_file(input_path, "DOCX")
    if source.suffix.lower() != ".docx":
        raise RenderError("render 输入文件必须为 .docx")
    if dpi <= 0:
        raise RenderError("dpi 必须为正整数")
    destination = Path(output_dir).expanduser().resolve()
    destination.mkdir(parents=True, exist_ok=True)
    for pattern in ("page-*.png", "vision-page-*.jpg"):
        for stale in destination.glob(pattern):
            stale.unlink()
    final_pdf = destination / f"{source.stem}.pdf"
    final_pdf.unlink(missing_ok=True)
    libreoffice = _which("libreoffice", "soffice")
    pdftoppm = _which("pdftoppm")
    with tempfile.TemporaryDirectory(prefix="oc-docx-lo-") as temporary:
        temp = Path(temporary)
        home = temp / "home"
        profile = temp / "profile"
        home.mkdir()
        profile.mkdir()
        env = os.environ.copy()
        env["HOME"] = str(home)
        command = [
            libreoffice,
            "--headless",
            "--nologo",
            "--nodefault",
            "--nolockcheck",
            "--nofirststartwizard",
            f"-env:UserInstallation={profile.resolve().as_uri()}",
            "--convert-to",
            "pdf:writer_pdf_Export",
            "--outdir",
            str(temp),
            str(source),
        ]
        result = subprocess.run(command, env=env, text=True, capture_output=True)
        if verbose:
            print(result.stdout, file=sys.stderr)
            print(result.stderr, file=sys.stderr)
        pdf = temp / f"{source.stem}.pdf"
        if result.returncode != 0 or not pdf.is_file() or pdf.stat().st_size == 0:
            raise RenderError(f"LibreOffice 转换失败：{result.stderr.strip() or result.stdout.strip()}")
        if emit_pdf:
            shutil.copy2(pdf, final_pdf)
        ppm = subprocess.run(
            [pdftoppm, "-png", "-r", str(dpi), str(pdf), str(destination / "page")],
            text=True,
            capture_output=True,
        )
        if verbose:
            print(ppm.stdout, file=sys.stderr)
            print(ppm.stderr, file=sys.stderr)
        if ppm.returncode != 0:
            raise RenderError(f"PDF 转 PNG 失败：{ppm.stderr.strip()}")
    pages = sorted(destination.glob("page-*.png"), key=_natural_page_number)
    if not pages:
        raise RenderError("没有生成页面 PNG")
    vision_pages = []
    for page in pages:
        number = _natural_page_number(page)
        vision_pages.append(_make_vision_copy(page, destination / f"vision-page-{number}.jpg"))
    return {
        "render_dir": str(destination),
        "pdf": str(final_pdf) if emit_pdf else None,
        "pages": [str(page) for page in pages],
        "vision_pages": [str(destination / item["file"]) for item in vision_pages],
    }


def _mostly_white_ratio(image_path: Path) -> float:
    _load_pillow_dependency()
    with Image.open(image_path) as image:
        histogram = image.convert("L").histogram()
        whiteish = sum(histogram[248:256])
        return round(whiteish / max(1, image.width * image.height), 6)


def inspect_docx(input_path: str | Path, render_dir: str | Path | None = None) -> dict[str, Any]:
    _load_docx_dependencies()
    source = _resolve_existing_file(input_path, "DOCX")
    with zipfile.ZipFile(source, "r") as archive:
        bad_member = archive.testzip()
        members = archive.namelist()
        image_members = [name for name in members if name.startswith("word/media/")]
        has_comments = "word/comments.xml" in members
        has_custom_properties = "docProps/custom.xml" in members
    document = Document(source)
    headings: list[dict[str, Any]] = []
    non_empty = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            non_empty += 1
        if paragraph.style and paragraph.style.name.startswith("Heading") and text:
            headings.append({"style": paragraph.style.name, "text": text})
    core = document.core_properties
    report: dict[str, Any] = {
        "file": str(source),
        "size_bytes": source.stat().st_size,
        "zip_integrity_ok": bad_member is None,
        "bad_zip_member": bad_member,
        "paragraph_count": len(document.paragraphs),
        "non_empty_paragraph_count": non_empty,
        "table_count": len(document.tables),
        "image_count": len(image_members),
        "headings": headings,
        "has_comments_xml": has_comments,
        "has_custom_properties": has_custom_properties,
        "core_properties": {
            "title": core.title or "",
            "subject": core.subject or "",
            "author": core.author or "",
            "last_modified_by": core.last_modified_by or "",
            "comments": core.comments or "",
            "keywords": core.keywords or "",
            "category": core.category or "",
        },
        "render_pairing_ok": True,
        "pages": [],
        "warnings": [],
    }
    if non_empty == 0 and not document.tables:
        report["warnings"].append("文档似乎没有可见内容")
    if has_comments:
        report["warnings"].append("文档包含批注 XML；交付前确认是否需要保留")
    if core.author or core.last_modified_by:
        report["warnings"].append("文档包含作者或最后修改人信息")
    if render_dir:
        _load_pillow_dependency()
        directory = Path(render_dir).expanduser().resolve()
        pages = sorted(directory.glob("page-*.png"), key=_natural_page_number)
        vision = { _natural_page_number(path): path for path in directory.glob("vision-page-*.jpg") }
        if not pages:
            report["warnings"].append("渲染目录中未找到 page-*.png")
            report["render_pairing_ok"] = False
        page_numbers = {_natural_page_number(path) for path in pages}
        if set(vision) != page_numbers:
            report["warnings"].append("原始页面与视觉质检副本未一一对应")
            report["render_pairing_ok"] = False
        for image_path in pages:
            number = _natural_page_number(image_path)
            with Image.open(image_path) as page_image:
                page: dict[str, Any] = {
                    "page": number,
                    "file": image_path.name,
                    "width": page_image.width,
                    "height": page_image.height,
                    "size_bytes": image_path.stat().st_size,
                    "mostly_white_ratio": _mostly_white_ratio(image_path),
                    "vision_copy": None,
                }
            vision_path = vision.get(number)
            if vision_path:
                with Image.open(vision_path) as vision_image:
                    page["vision_copy"] = {
                        "file": vision_path.name,
                        "width": vision_image.width,
                        "height": vision_image.height,
                        "size_bytes": vision_path.stat().st_size,
                    }
                if vision_path.stat().st_size >= VISION_COPY_MAX_BYTES:
                    report["warnings"].append(f"{vision_path.name} 超过视觉质检大小预算")
                    report["render_pairing_ok"] = False
            report["pages"].append(page)
            if page["mostly_white_ratio"] > 0.995:
                report["warnings"].append(f"{image_path.name} 可能接近空白页")
            if page["width"] < 800 or page["height"] < 1000:
                report["warnings"].append(f"{image_path.name} 分辨率偏低")
    return report


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
CP_NS = "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
DC_NS = "http://purl.org/dc/elements/1.1/"
DCT_NS = "http://purl.org/dc/terms/"
EP_NS = "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CUSTOM_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/custom-properties"


def _remove_nodes(root: Any, tags: list[str]) -> None:
    for tag in tags:
        for node in list(root.findall(tag)):
            root.remove(node)


def _scrub_xml(name: str, data: bytes, keep_title: bool, keep_author: bool) -> bytes:
    _load_lxml_dependency()
    root = etree.fromstring(data, etree.XMLParser(remove_blank_text=False))
    if name == "docProps/core.xml":
        removable = [
            f"{{{CP_NS}}}lastModifiedBy",
            f"{{{DC_NS}}}subject",
            f"{{{CP_NS}}}keywords",
            f"{{{DC_NS}}}description",
            f"{{{CP_NS}}}category",
            f"{{{DCT_NS}}}created",
            f"{{{DCT_NS}}}modified",
            f"{{{CP_NS}}}lastPrinted",
            f"{{{CP_NS}}}revision",
        ]
        if not keep_title:
            removable.append(f"{{{DC_NS}}}title")
        if not keep_author:
            removable.append(f"{{{DC_NS}}}creator")
        _remove_nodes(root, removable)
    elif name == "docProps/app.xml":
        for tag_name in ("Company", "Manager", "Template", "HyperlinkBase"):
            node = root.find(f"{{{EP_NS}}}{tag_name}")
            if node is not None:
                node.text = ""
    elif name == "[Content_Types].xml":
        for node in list(root.findall(f"{{{CT_NS}}}Override")):
            if node.get("PartName") == "/docProps/custom.xml":
                root.remove(node)
    elif name == "_rels/.rels":
        for node in list(root.findall(f"{{{REL_NS}}}Relationship")):
            if node.get("Type") == CUSTOM_REL_TYPE or node.get("Target") == "docProps/custom.xml":
                root.remove(node)
    if name.startswith("word/") and name.endswith(".xml"):
        for element in root.iter():
            for attribute in list(element.attrib):
                if attribute.startswith(f"{{{W_NS}}}rsid"):
                    del element.attrib[attribute]
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=None)


def scrub_docx(
    input_path: str | Path,
    output_path: str | Path,
    *,
    keep_title: bool = False,
    keep_author: bool = False,
) -> Path:
    source = _resolve_existing_file(input_path, "DOCX")
    destination = _resolve_output(output_path)
    if source == destination:
        raise ValueError("scrub 输出文件不能覆盖输入文件")
    with tempfile.TemporaryDirectory(prefix="oc-docx-scrub-") as temporary:
        temporary_output = Path(temporary) / destination.name
        with zipfile.ZipFile(source, "r") as incoming, zipfile.ZipFile(temporary_output, "w", zipfile.ZIP_DEFLATED) as outgoing:
            for info in incoming.infolist():
                if info.filename == "docProps/custom.xml":
                    continue
                data = incoming.read(info.filename)
                if info.filename.endswith((".xml", ".rels")) or info.filename == "[Content_Types].xml":
                    try:
                        data = _scrub_xml(info.filename, data, keep_title, keep_author)
                    except etree.XMLSyntaxError:
                        pass
                outgoing.writestr(info, data)
        with zipfile.ZipFile(temporary_output, "r") as check:
            bad_member = check.testzip()
            if bad_member:
                raise RuntimeError(f"scrub 后 DOCX ZIP 损坏：{bad_member}")
        shutil.copy2(temporary_output, destination)
    return destination


def _build_command(args: argparse.Namespace) -> int:
    config = apply_overrides(load_config(args.input), args.set)
    output = DocumentBuilder(config).build(
        args.output,
        cover=False if args.no_cover else None,
        toc=False if args.no_toc else None,
    )
    print(output)
    return 0


def _convert_command(args: argparse.Namespace) -> int:
    engine = "quarto" if args.quarto else "pandoc" if args.pandoc else None
    output = convert_docx(
        args.input,
        args.output,
        engine=engine,
        reference_doc=args.reference_doc,
        use_reference=not args.no_reference_doc,
    )
    print(output)
    return 0


def _render_command(args: argparse.Namespace) -> int:
    result = render_docx(args.input, args.output_dir, emit_pdf=args.emit_pdf, dpi=args.dpi, verbose=args.verbose)
    print(json.dumps(result, ensure_ascii=False))
    return 0


def _inspect_command(args: argparse.Namespace) -> int:
    report = inspect_docx(args.input, args.render_dir)
    if args.json:
        output = _resolve_output(args.json)
        output.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        report["report_file"] = str(output)
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0 if report["zip_integrity_ok"] and report["render_pairing_ok"] else 1


def _scrub_command(args: argparse.Namespace) -> int:
    print(scrub_docx(args.input, args.output, keep_title=args.keep_title, keep_author=args.keep_author))
    return 0


def _add_convert_arguments(parser: argparse.ArgumentParser) -> None:
    engines = parser.add_mutually_exclusive_group()
    engines.add_argument("--quarto", action="store_true", help="使用 Quarto")
    engines.add_argument("--pandoc", action="store_true", help="使用 Pandoc")
    parser.add_argument("--reference-doc", help="自定义 Word reference.docx")
    parser.add_argument("--no-reference-doc", action="store_true", help="不使用默认 reference.docx")
    parser.add_argument("-o", "--output", help="输出 DOCX；默认与输入同目录同名")
    parser.add_argument("input", help="输入 .md/.markdown/.qmd")


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        prog="oc-docx",
        description="高质量 Word DOCX 生成、逐页渲染、检查和隐私清理",
    )
    subparsers = parser.add_subparsers(dest="command", required=True)
    convert = subparsers.add_parser(
        "convert",
        help="从 Markdown/Quarto 生成 DOCX（原生 OMML 公式）",
        description="从 Markdown/Quarto 生成 DOCX（原生 OMML 公式）",
    )
    _add_convert_arguments(convert)
    convert.set_defaults(func=_convert_command)
    build = subparsers.add_parser(
        "build",
        help="从 YAML/JSON 结构化配置生成 DOCX",
        description="从 YAML/JSON 结构化配置生成 DOCX",
    )
    build.add_argument("input", help="YAML 或 JSON 配置")
    build.add_argument("-o", "--output", required=True, help="输出 DOCX")
    build.add_argument("--no-cover", action="store_true", help="忽略配置中的封面")
    build.add_argument("--no-toc", action="store_true", help="忽略配置中的目录")
    build.add_argument("--set", action="append", default=[], help="覆盖配置 key=value，可重复")
    build.set_defaults(func=_build_command)
    render = subparsers.add_parser(
        "render",
        help="DOCX → PDF + 原始逐页 PNG + 视觉安全 JPEG",
        description="DOCX → PDF + 原始逐页 PNG + 视觉安全 JPEG",
    )
    render.add_argument("input", help="输入 DOCX")
    render.add_argument("-o", "--output-dir", required=True, help="输出目录")
    render.add_argument("--emit-pdf", action="store_true", help="保留中间 PDF")
    render.add_argument("--dpi", type=int, default=160, help="原始 PNG 分辨率，默认 160")
    render.add_argument("--verbose", action="store_true", help="显示渲染器日志")
    render.set_defaults(func=_render_command)
    inspect = subparsers.add_parser(
        "inspect",
        help="检查 DOCX 结构、元数据和逐页配对",
        description="检查 DOCX 结构、元数据和逐页配对",
    )
    inspect.add_argument("input", help="输入 DOCX")
    inspect.add_argument("--render-dir", help="包含 page-*.png/vision-page-*.jpg 的目录")
    inspect.add_argument("--json", help="保存检查报告 JSON")
    inspect.set_defaults(func=_inspect_command)
    scrub = subparsers.add_parser(
        "scrub",
        help="清理 DOCX 隐私元数据和 rsid",
        description="清理 DOCX 隐私元数据和 rsid",
    )
    scrub.add_argument("input", help="输入 DOCX")
    scrub.add_argument("-o", "--output", required=True, help="输出 DOCX")
    scrub.add_argument("--keep-title", action="store_true", help="保留文档属性中的标题")
    scrub.add_argument("--keep-author", action="store_true", help="保留用户明确要求的作者")
    scrub.set_defaults(func=_scrub_command)
    return parser


def main(argv: list[str] | None = None) -> int:
    actual = list(sys.argv[1:] if argv is None else argv)
    commands = {"convert", "build", "render", "inspect", "scrub"}
    if actual and actual[0] not in commands and actual[0] not in {"-h", "--help"}:
        actual.insert(0, "convert")
    parser = build_parser()
    if not actual:
        parser.print_help(sys.stderr)
        return 2
    args = parser.parse_args(actual)
    try:
        return int(args.func(args))
    except (ConfigError, RenderError, FileNotFoundError, RuntimeError, ValueError, subprocess.CalledProcessError, zipfile.BadZipFile) as exc:
        print(f"oc-docx: {exc}", file=sys.stderr)
        return 2


if __name__ == "__main__":
    raise SystemExit(main())
